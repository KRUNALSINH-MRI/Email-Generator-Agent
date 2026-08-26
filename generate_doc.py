"""Generate the POC Documentation Word file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else None
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Email Generator Agent — POC Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    add_paragraph(doc, "Project: Email Generator Agent", bold=True)
    add_paragraph(doc, "Author: Krunalsinh Chhasatiya")
    add_paragraph(doc, "Date: August 2026")
    add_paragraph(doc, "Version: 2.0 (Full Coverage — Tasks 1 & 2)")
    doc.add_paragraph("")

    # ─── 1. Overview ───
    add_heading(doc, "1. Overview")
    doc.add_paragraph(
        "The Email Generator Agent is an AI-powered professional email generator "
        "built with LangGraph, Azure OpenAI, and MCP (Model Context Protocol). "
        "It features:"
    )
    add_bullet(doc, "A LangGraph-based agent that validates inputs, fetches tone-specific guidelines via MCP, and generates structured emails using Azure OpenAI.")
    add_bullet(doc, "An MCP App with a React + TypeScript UI for interactive email generation, review, editing, and approval/rejection.")
    add_bullet(doc, "A headless hosted agent deployable to Azure AI Foundry with the Invocations protocol.")
    add_bullet(doc, "Docker containerization and Azure Bicep infrastructure-as-code for cloud deployment.")

    # ─── 2. Architecture ───
    add_heading(doc, "2. Architecture")

    add_heading(doc, "2.1 High-Level Flow", level=2)
    doc.add_paragraph(
        "The system operates in two modes:"
    )
    add_paragraph(doc, "Mode 1: MCP App (Interactive UI)", bold=True)
    doc.add_paragraph(
        "User → React UI → REST API (http_server.py) → MCP Tools → "
        "email_service.py → LangGraph Agent → Azure OpenAI → Response → UI"
    )
    add_paragraph(doc, "Mode 2: Hosted Agent (Headless API)", bold=True)
    doc.add_paragraph(
        "Client → POST /invocations → hosted_agent.py → LangGraph Agent → "
        "Azure OpenAI → JSON Response"
    )

    add_heading(doc, "2.2 Component Diagram", level=2)
    add_code_block(doc, """┌─────────────────────────────────────────────────────────────┐
│                    MCP App (Port 8000)                       │
│  ┌──────────┐    ┌──────────────┐    ┌────────────────┐    │
│  │ React UI │───►│ REST Bridge  │───►│  MCP Tools     │    │
│  │ (Vite)   │    │ /api/*       │    │ generate_email │    │
│  └──────────┘    └──────────────┘    │ approve_email  │    │
│                                       │ reject_email   │    │
│                                       └───────┬────────┘    │
└───────────────────────────────────────────────┼─────────────┘
                                                │
                    ┌───────────────────────────▼──────────────┐
                    │         LangGraph Agent (graph.py)        │
                    │  validate → get_guidelines → prompt → LLM│
                    └───────────────────┬──────────────────────┘
                                        │
              ┌─────────────────────────▼────────────────────┐
              │                                               │
    ┌─────────▼──────────┐              ┌────────────────────▼─┐
    │  MCP Guidelines    │              │   Azure OpenAI       │
    │  Server (stdio)    │              │   (GPT Model)        │
    │  server.py         │              └──────────────────────┘
    └────────────────────┘""")

    # ─── 3. Project Structure ───
    add_heading(doc, "3. Project Structure")
    add_code_block(doc, """├── src/email_generator/
│   ├── graph.py            # LangGraph workflow definition
│   ├── hosted_agent.py     # Azure Foundry host server
│   ├── llm.py              # Azure OpenAI client setup
│   ├── prompts.py          # Email generation prompt template
│   ├── state.py            # State and output schemas
│   ├── mcp_client.py       # MCP client (stdio or remote)
│   └── main.py             # CLI entry point
├── mcp_server/
│   ├── server.py           # FastMCP guidelines server (stdio)
│   ├── http_server.py      # MCP App: HTTP server + REST bridge + UI
│   ├── email_service.py    # Agent invocation wrapper
│   └── ui/                 # React + TypeScript MCP App UI
│       ├── src/App.tsx     # Main app component
│       ├── package.json    # Node dependencies
│       └── vite.config.ts  # Vite bundler config
├── tests/
│   ├── test_agent.py           # Agent unit tests
│   ├── test_deployed_agent.py  # Integration tests
│   ├── test_mcp.py             # MCP server tests
│   └── test_mcp_http.py        # MCP HTTP transport tests
├── infra/                  # Azure Bicep infrastructure
│   ├── main.bicep
│   ├── main.json
│   └── main.parameters.json
├── azure.yaml              # Azure Developer CLI configuration
├── Dockerfile              # Container image for MCP App
├── pyproject.toml          # Python project configuration
├── start_agent.py          # Local hosted agent entry point
└── requirements.txt        # Python dependencies""")

    # ─── 4. Tech Stack ───
    add_heading(doc, "4. Technology Stack")
    add_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["AI/LLM", "Azure OpenAI (GPT)", "Email text generation with structured output"],
            ["Agent Framework", "LangGraph", "Stateful workflow orchestration"],
            ["Tool Protocol", "MCP (FastMCP)", "Tool discovery and invocation"],
            ["Backend", "Starlette + Uvicorn", "Async HTTP server"],
            ["Frontend", "React + TypeScript", "Interactive email UI"],
            ["Bundler", "Vite", "Fast frontend build and HMR"],
            ["Hosting", "Azure AI Foundry", "Hosted agent with Invocations protocol"],
            ["Infrastructure", "Azure Bicep", "Infrastructure-as-Code"],
            ["Deployment", "Azure Developer CLI (azd)", "Provision + deploy workflow"],
            ["Container", "Docker", "Containerized deployment"],
            ["Package Manager", "uv", "Fast Python dependency management"],
        ],
    )

    # ─── 5. Task 1: Email Generator Agent ───
    add_heading(doc, "5. Task 1: Email Generator Agent")

    add_heading(doc, "5.1 LangGraph Workflow", level=2)
    doc.add_paragraph("The agent follows a sequential graph:")
    add_code_block(doc, "START → validate_input → [error | continue] → get_mcp_guidelines → build_prompt → generate_email → END")

    add_heading(doc, "5.2 Graph Nodes", level=2)
    add_table(doc,
        ["Node", "Description"],
        [
            ["validate_input", "Checks tone, context, and data_points are present"],
            ["get_mcp_guidelines", "Calls MCP server for tone-specific writing guidelines"],
            ["build_prompt", "Constructs the LLM prompt with context + guidelines"],
            ["generate_email", "Calls Azure OpenAI with structured output (subject + email)"],
        ],
    )

    add_heading(doc, "5.3 MCP Guidelines Server (server.py)", level=2)
    doc.add_paragraph(
        "A lightweight FastMCP server providing the get_email_guidelines tool. "
        "The LangGraph agent connects to it via stdio transport. Supports four tones:"
    )
    add_bullet(doc, "formal — Professional language, clear structure, no casual expressions")
    add_bullet(doc, "friendly — Warm, approachable, conversational professional language")
    add_bullet(doc, "empathetic — Acknowledges concerns, shows understanding, reassuring")
    add_bullet(doc, "assertive — Direct, confident, clear action without aggression")

    add_heading(doc, "5.4 MCP Client (mcp_client.py)", level=2)
    doc.add_paragraph(
        "Supports two connection modes:"
    )
    add_bullet(doc, "Local (default): Spawns server.py as a subprocess via stdio transport")
    add_bullet(doc, "Remote: Connects to MCP_SERVER_URL environment variable (HTTP transport)")

    add_heading(doc, "5.5 Hosted Agent (hosted_agent.py)", level=2)
    doc.add_paragraph(
        "Implements the Azure AI Foundry Invocations protocol. Exposes a POST /invocations "
        "endpoint on port 8088 that accepts JSON input and returns the generated email."
    )

    add_heading(doc, "5.6 Input / Output", level=2)
    add_paragraph(doc, "Request:", bold=True)
    add_code_block(doc, """{
  "tone": "formal",
  "context": "Inform the client that the project deployment has been completed.",
  "data_points": [
    "Deployment completed on August 25, 2026",
    "All planned features were deployed",
    "Initial validation testing passed",
    "No critical issues were found"
  ]
}""")
    add_paragraph(doc, "Response:", bold=True)
    add_code_block(doc, """{
  "subject": "Project Deployment Completed Successfully",
  "email": "Dear Client,\\n\\nWe are pleased to inform you that..."
}""")

    # ─── 6. Task 2: MCP App (UI) ───
    add_heading(doc, "6. Task 2: MCP App with Interactive UI")

    add_heading(doc, "6.1 Overview", level=2)
    doc.add_paragraph(
        "A full-stack MCP App providing a React-based web interface for users to "
        "interactively generate, review, modify, and approve/reject emails. "
        "The UI communicates with MCP tools through a REST API bridge."
    )

    add_heading(doc, "6.2 MCP Tools (http_server.py)", level=2)
    add_table(doc,
        ["Tool", "Parameters", "Description"],
        [
            ["get_email_guidelines", "tone: str", "Returns tone-specific writing guidelines"],
            ["generate_email", "tone, context, data_points", "Generates email via LangGraph agent"],
            ["approve_email", "subject, email", "Approves and sends the email"],
            ["reject_email", "subject, email", "Rejects the generated email"],
        ],
    )

    add_heading(doc, "6.3 REST API Bridge", level=2)
    doc.add_paragraph(
        "The React UI cannot call MCP tools directly. A REST bridge translates "
        "HTTP requests into MCP tool calls:"
    )
    add_table(doc,
        ["Endpoint", "Method", "MCP Tool Called"],
        [
            ["/api/generate-email", "POST", "generate_email"],
            ["/api/approve-email", "POST", "approve_email"],
            ["/api/reject-email", "POST", "reject_email"],
        ],
    )

    add_heading(doc, "6.4 React UI Features", level=2)
    add_bullet(doc, "Tone selector dropdown (formal, friendly, empathetic, assertive)")
    add_bullet(doc, "Context textarea for email purpose/background")
    add_bullet(doc, "Data points input (one per line)")
    add_bullet(doc, "Generate Email button with loading spinner")
    add_bullet(doc, "Editable subject and email body fields after generation")
    add_bullet(doc, "Approve & Send button — calls approve_email tool, shows confirmation")
    add_bullet(doc, "Reject button — calls reject_email tool, shows rejection message")
    add_bullet(doc, "Error display for failed operations")

    add_heading(doc, "6.5 UI Tech Stack", level=2)
    add_bullet(doc, "React 19 with TypeScript")
    add_bullet(doc, "Vite 8 for bundling and development")
    add_bullet(doc, "Built output served from mcp_server/ui/dist/ by Starlette StaticFiles")

    add_heading(doc, "6.6 User Flow", level=2)
    doc.add_paragraph("1. User opens http://localhost:8000")
    doc.add_paragraph("2. Fills in tone, context, and data points")
    doc.add_paragraph("3. Clicks 'Generate Email' → calls /api/generate-email → MCP tool → agent")
    doc.add_paragraph("4. Generated email appears with editable subject and body")
    doc.add_paragraph("5. User can modify the email text directly")
    doc.add_paragraph("6. Clicks 'Approve & Send' → calls /api/approve-email → returns 'Email has been sent'")
    doc.add_paragraph("7. Or clicks 'Reject' → calls /api/reject-email → returns 'Email has been rejected'")

    # ─── 7. Deployment ───
    add_heading(doc, "7. Deployment")

    add_heading(doc, "7.1 Azure Developer CLI", level=2)
    add_code_block(doc, """az login
azd up        # Provision infrastructure + deploy
azd deploy    # Deploy only (after initial setup)""")

    add_heading(doc, "7.2 Docker", level=2)
    add_code_block(doc, """docker build -t email-generator .
docker run -p 8000:8000 \\
  -e AZURE_OPENAI_ENDPOINT=<endpoint> \\
  -e AZURE_AI_MODEL_DEPLOYMENT_NAME=<model> \\
  email-generator""")

    add_heading(doc, "7.3 Infrastructure (Bicep)", level=2)
    doc.add_paragraph(
        "Azure Bicep templates in the infra/ directory provision the required "
        "Azure resources: AI Foundry project, OpenAI resource, and container hosting."
    )

    # ─── 8. Running Locally ───
    add_heading(doc, "8. Running Locally")

    add_heading(doc, "8.1 Prerequisites", level=2)
    add_bullet(doc, "Python 3.14+")
    add_bullet(doc, "Node.js 18+ (for UI)")
    add_bullet(doc, "uv (Python package manager)")
    add_bullet(doc, "Azure CLI & Azure Developer CLI")
    add_bullet(doc, "Azure OpenAI resource with deployed model")

    add_heading(doc, "8.2 Setup", level=2)
    add_code_block(doc, """git clone https://github.com/KRUNALSINH-MRI/Email-Generator-Agent.git
cd Email-Generator-Agent
uv venv && source .venv/bin/activate
uv sync
cd mcp_server/ui && npm install && cd ../..
cp .env.example .env
# Set AZURE_OPENAI_ENDPOINT and AZURE_AI_MODEL_DEPLOYMENT_NAME""")

    add_heading(doc, "8.3 Run MCP App (UI + Server)", level=2)
    add_code_block(doc, """cd mcp_server/ui && npm run build && cd ../..
PYTHONPATH=src python -m mcp_server.http_server
# Open http://localhost:8000""")

    add_heading(doc, "8.4 Run Hosted Agent (Headless)", level=2)
    add_code_block(doc, """python start_agent.py
# Test: curl -s -X POST http://localhost:8088/invocations ...""")

    add_heading(doc, "8.5 UI Development (Hot Reload)", level=2)
    add_code_block(doc, """# Terminal 1:
PYTHONPATH=src python -m mcp_server.http_server
# Terminal 2:
cd mcp_server/ui && npm run dev
# Open http://localhost:5173""")

    # ─── 9. Testing ───
    add_heading(doc, "9. Testing")
    add_code_block(doc, """pytest tests/ -v                    # All tests
pytest tests/test_agent.py -v        # Agent unit tests
pytest tests/test_mcp.py -v          # MCP server tests
pytest tests/test_mcp_http.py -v     # HTTP transport tests
pytest tests/test_deployed_agent.py -v  # Deployed agent tests""")

    # ─── 10. Environment Variables ───
    add_heading(doc, "10. Environment Variables")
    add_table(doc,
        ["Variable", "Description", "Required"],
        [
            ["AZURE_OPENAI_ENDPOINT", "Azure OpenAI resource endpoint", "Yes"],
            ["AZURE_AI_MODEL_DEPLOYMENT_NAME", "Deployed model name", "Yes"],
            ["MCP_SERVER_URL", "Remote MCP server URL (overrides local stdio)", "No"],
        ],
    )

    # ─── 11. Summary ───
    add_heading(doc, "11. Summary")
    add_table(doc,
        ["Task", "Deliverable", "Status"],
        [
            ["Task 1", "LangGraph Email Agent + MCP Guidelines Server + Azure Foundry Hosting", "Complete"],
            ["Task 2", "MCP App with React UI + MCP Tools + REST Bridge + Docker", "Complete"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "The project demonstrates a production-ready AI agent architecture combining "
        "LangGraph for workflow orchestration, MCP for tool interoperability, "
        "Azure OpenAI for generation, and a modern React frontend for user interaction."
    )

    # Save
    output_path = "Email_Generator_Agent_POC_Documentation.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    main()
